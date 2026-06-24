#!/usr/bin/env python3
"""Build 5 eye-catchy files (3 docx + 2 xlsx) for upload to Google Drive (convert to Docs/Sheets)."""
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DIR="/Users/user/Claude/eBook/kit-resources/"
# palette
NAVY="14182B"; TEAL="00C896"; BLUE="5B8CFF"; INK="1A1F36"; GREY="6B7299"
MINT="E8FBF4"; BLUETINT="EEF3FF"; YELLOW="FFF6D6"; LINE="DDE3F0"

# ---------------- DOCX helpers ----------------
def base(doc):
    st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(11); st.font.color.rgb=RGBColor.from_string(INK)
    sec=doc.sections[0]
    for m in ("top_margin","bottom_margin"): setattr(sec,m,Inches(0.6))
    for m in ("left_margin","right_margin"): setattr(sec,m,Inches(0.8))

def shade(cell,color):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),color); tcPr.append(shd)

def no_borders(tbl):
    tblPr=tbl._tbl.tblPr; b=OxmlElement('w:tblBorders')
    for e in ("top","left","bottom","right","insideH","insideV"):
        x=OxmlElement('w:'+e); x.set(qn('w:val'),'none'); b.append(x)
    tblPr.append(b)

def banner(doc,title,sub):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width=Inches(6.9); no_borders(t)
    c=t.rows[0].cells[0]; shade(c,NAVY)
    c.paragraphs[0].paragraph_format.space_before=Pt(10)
    r=c.paragraphs[0].add_run(title); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=RGBColor.from_string("FFFFFF")
    if sub:
        p=c.add_paragraph(); p.paragraph_format.space_after=Pt(10)
        rr=p.add_run(sub); rr.font.size=Pt(10.5); rr.font.color.rgb=RGBColor.from_string(TEAL)
    doc.add_paragraph()

def section(doc,text,color=TEAL):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(2)
    r=p.add_run("▍ "+text); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string(color)

def chip(doc,text,fill,txtcolor=INK):
    t=doc.add_table(rows=1,cols=1); t.columns[0].width=Inches(6.9); no_borders(t)
    c=t.rows[0].cells[0]; shade(c,fill)
    r=c.paragraphs[0].add_run(text); r.font.size=Pt(10.5); r.font.color.rgb=RGBColor.from_string(txtcolor)
    c.paragraphs[0].paragraph_format.space_before=Pt(4); c.paragraphs[0].paragraph_format.space_after=Pt(4)

def styled_table(doc,headers,rows,hfill=NAVY,accent_last=False):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; shade(c,hfill)
        r=c.paragraphs[0].add_run(h); r.bold=True; r.font.color.rgb=RGBColor.from_string("FFFFFF"); r.font.size=Pt(10.5)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        last=accent_last and ri==len(rows)-1
        for i,v in enumerate(row):
            if last: shade(cells[i],MINT)
            else: shade(cells[i], "FFFFFF" if ri%2==0 else "F5F7FC")
            run=cells[i].paragraphs[0].add_run(str(v)); run.font.size=Pt(10.5)
            if last: run.bold=True; run.font.color.rgb=RGBColor.from_string(INK)
    return t

def foot(doc):
    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("NextGen AI Lab  ·  Powered by OptimityFX  ·  support@optimityfx.com")
    r.italic=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(GREY)

def bullets(doc,items,color=TEAL):
    for it in items:
        p=doc.add_paragraph(style="List Bullet");
        r=p.add_run(it); r.font.size=Pt(11)

# ============ 1) PORTFOLIO ============
doc=Document(); base(doc)
banner(doc,"[Your Name] — [Your Service]","Helping [audience] get [result] using AI-powered systems")
chip(doc,"📧 [your@email.com]      📱 [WhatsApp]      🔗 [Instagram / LinkedIn]",BLUETINT)
section(doc,"About Me")
doc.add_paragraph("I help [type of business] with [your service]. I use AI tools to deliver fast, high-quality work at a fair price. [2–3 honest lines — mention any relevant background, even if you're new.]")
section(doc,"What I Do")
bullets(doc,["[Service 1]","[Service 2]","[Service 3]"])
section(doc,"Work Samples",BLUE)
doc.add_paragraph("Replace with your 3 best samples (screenshots/links). New? Make 3 free samples first — that IS your portfolio.")
for i in (1,2,3):
    p=doc.add_paragraph(); r=p.add_run(f"Sample {i} — [Client / Project]"); r.bold=True; r.font.color.rgb=RGBColor.from_string(INK)
    doc.add_paragraph("What I did: [1 line]   ·   Result: [1 line]")
section(doc,"Packages")
styled_table(doc,["Package","What's included","Price"],
    [["Starter","[deliverables]","₹[xxx]"],["Growth","[deliverables]","₹[xxx]"],["Pro","[deliverables]","₹[xxx]"]])
section(doc,"Let's Work Together",BLUE)
p=doc.add_paragraph(); r=p.add_run("Message me on [WhatsApp/Instagram] or email [your@email.com]."); r.bold=True
foot(doc); doc.save(DIR+"01-Portfolio-Template.docx"); print("✓ Portfolio")

# ============ 2) SERVICE MENU ============
doc=Document(); base(doc)
banner(doc,"[Your Name / Brand] — Service Menu","AI-powered [service] for [target clients] · Fast turnaround, fair pricing")
chip(doc,"📧 [your@email.com]      📱 [WhatsApp]      🔗 [social handle]",BLUETINT)
for nm,price,who,items,clr in [
 ("🟢 Starter Package","₹[price]","businesses just getting started",["[Deliverable 1]","[Deliverable 2]","[Deliverable 3]","Turnaround: [X days]"],TEAL),
 ("🔵 Growth Package — MOST POPULAR","₹[price] / month","businesses that want consistent results",["Everything in Starter, plus:","[Deliverable 4]","[Deliverable 5]","Monthly report + priority support"],BLUE),
 ("🟣 Pro Package","₹[price] / month","businesses ready to scale",["Everything in Growth, plus:","[Deliverable 6]","[Deliverable 7]","[Add-on]"],"7A5BFF")]:
    section(doc,nm,clr)
    p=doc.add_paragraph(); r=p.add_run(f"{price}   ·   Perfect for: {who}"); r.bold=True; r.font.color.rgb=RGBColor.from_string(INK)
    bullets(doc,items)
section(doc,"Add-Ons")
styled_table(doc,["Add-on","Price"],[["[Extra service 1]","₹[xxx]"],["[Extra service 2]","₹[xxx]"],["Rush delivery (24h)","+[xx]%"]])
section(doc,"How It Works",BLUE)
bullets(doc,["You pick a package.","50% advance to start (UPI / bank).","I deliver within the timeline.","50% on delivery. Done!"])
chip(doc,"💡 Tip: keep prices ending in 9 (₹2,999) and highlight the middle option — most clients pick it.",YELLOW)
foot(doc); doc.save(DIR+"02-Service-Menu.docx"); print("✓ Service Menu")

# ============ 3) PROPOSAL ============
doc=Document(); base(doc)
banner(doc,"Project Proposal","Prepared by [Your Name / Brand]")
styled_table(doc,["Prepared for","Date"],[["[Client Name / Business]","[Date]"]])
section(doc,"1. What You Need")
doc.add_paragraph("[Describe the client's problem in 1–2 lines. Show you understood them.]")
section(doc,"2. What I'll Deliver")
bullets(doc,["[Deliverable 1 — with quantity]","[Deliverable 2]","[Deliverable 3]","[Deliverable 4]"])
section(doc,"3. Timeline",BLUE)
styled_table(doc,["Phase","What happens","When"],
    [["Kickoff","Brief + access + samples","Day 1"],["Draft","First batch for review","Day [x]"],["Delivery","Final, revised","Day [x]"]])
section(doc,"4. Investment")
styled_table(doc,["Item","Price"],[["[Package name]","₹[price]"],["[Add-on, optional]","₹[price]"],["TOTAL","₹[total]"]],accent_last=True)
chip(doc,"Payment terms: 50% advance to begin, 50% on delivery (UPI / bank transfer).",MINT)
section(doc,"5. Why Me",BLUE)
bullets(doc,["[Reason 1 — e.g. fast AI-assisted turnaround]","[Reason 2 — link 2 samples]","[Reason 3]"])
chip(doc,'✅ Ready to start? Reply "yes" and I\'ll send the payment link. Valid until [date].',BLUETINT,INK)
foot(doc); doc.save(DIR+"05-Proposal-Template.docx"); print("✓ Proposal")

# ---------------- XLSX helpers ----------------
navy=PatternFill("solid",fgColor=NAVY); teal=PatternFill("solid",fgColor=TEAL)
mint=PatternFill("solid",fgColor=MINT); yel=PatternFill("solid",fgColor=YELLOW)
zebra=PatternFill("solid",fgColor="F5F7FC"); bluef=PatternFill("solid",fgColor=BLUE)
white=Font(color="FFFFFF",bold=True); boldink=Font(bold=True,color=INK)
thin=Side(style="thin",color=LINE); bd=Border(left=thin,right=thin,top=thin,bottom=thin)
ctr=Alignment(horizontal="center",vertical="center"); wrap=Alignment(wrap_text=True,vertical="center",horizontal="center")

# ============ 4) INVOICE ============
wb=Workbook(); ws=wb.active; ws.title="Invoice"; ws.sheet_view.showGridLines=False
for i,w in enumerate([3,28,10,14,16],1): ws.column_dimensions[get_column_letter(i)].width=w
ws.merge_cells("B2:E2"); t=ws["B2"]; t.value="INVOICE"; t.font=Font(bold=True,size=26,color="FFFFFF"); t.fill=navy; t.alignment=Alignment(horizontal="left",vertical="center"); ws.row_dimensions[2].height=42
ws.merge_cells("B3:E3"); s=ws["B3"]; s.value="NextGen AI Lab · Powered by OptimityFX"; s.font=Font(color=TEAL,bold=True,size=10); s.fill=navy
ws["B5"]="[Your Name / Brand]"; ws["B5"].font=boldink
ws["B6"]="[Address / City]"; ws["B7"]="[your@email.com] · [phone]"; ws["B8"]="GSTIN: [if applicable]"
for lab,val,r in [("Invoice #","[INV-001]",5),("Date","[DD/MM/YYYY]",6),("Due date","[DD/MM/YYYY]",7)]:
    ws.cell(row=r,column=4,value=lab).font=boldink; ws.cell(row=r,column=4).alignment=Alignment(horizontal="right")
    c=ws.cell(row=r,column=5,value=val); c.fill=yel; c.border=bd; c.alignment=ctr
ws["B10"]="BILL TO"; ws["B10"].font=Font(bold=True,color=TEAL,size=11)
for r in (11,12):
    c=ws.cell(row=r,column=2,value="[Client name / business]" if r==11 else "[Client email / phone]");
    ws.merge_cells(start_row=r,start_column=2,end_row=r,end_column=3); c.fill=yel; c.border=bd
hr=14
for i,h in enumerate(["#","Description","Qty","Rate (₹)","Amount (₹)"],1):
    c=ws.cell(row=hr,column=i,value=h); c.fill=navy; c.font=white; c.alignment=ctr; c.border=bd
ws.row_dimensions[hr].height=22
first=hr+1; last=first+5
for r in range(first,last):
    ws.cell(row=r,column=1,value=r-hr).alignment=ctr; ws.cell(row=r,column=1).border=bd
    d=ws.cell(row=r,column=2,value="[Service description]"); d.fill=yel; d.border=bd
    q=ws.cell(row=r,column=3,value=1); q.fill=yel; q.border=bd; q.alignment=ctr
    rt=ws.cell(row=r,column=4,value=0 if r==first else None); rt.fill=yel; rt.border=bd; rt.number_format='#,##0'
    am=ws.cell(row=r,column=5,value=f"=C{r}*D{r}"); am.border=bd; am.number_format='#,##0'
    if r%2==0:
        for col in (1,5): ws.cell(row=r,column=col).fill=zebra
sub=last
ws.cell(row=sub,column=4,value="Subtotal").font=boldink; ws.cell(row=sub,column=4).alignment=Alignment(horizontal="right")
sc=ws.cell(row=sub,column=5,value=f"=SUM(E{first}:E{last-1})"); sc.number_format='#,##0'; sc.font=boldink; sc.border=bd
g=sub+1
ws.cell(row=g,column=2,value="GST % (set 0 if none)").font=Font(italic=True,size=9,color=GREY)
gc=ws.cell(row=g,column=3,value=18); gc.fill=yel; gc.border=bd; gc.alignment=ctr
ws.cell(row=g,column=4,value="GST amount").font=boldink; ws.cell(row=g,column=4).alignment=Alignment(horizontal="right")
ga=ws.cell(row=g,column=5,value=f"=E{sub}*C{g}/100"); ga.number_format='#,##0'; ga.border=bd
tot=g+1
tl=ws.cell(row=tot,column=4,value="TOTAL"); tl.fill=teal; tl.font=Font(bold=True,color=INK,size=12); tl.alignment=Alignment(horizontal="right",vertical="center")
tv=ws.cell(row=tot,column=5,value=f"=E{sub}+E{g}"); tv.fill=teal; tv.font=Font(bold=True,color=INK,size=12); tv.number_format='#,##0'; tv.border=bd
ws.row_dimensions[tot].height=24
p=tot+2
ws.cell(row=p,column=2,value="PAYMENT DETAILS").font=Font(bold=True,color=TEAL,size=11)
ws.cell(row=p+1,column=2,value="UPI: [your-upi@bank]")
ws.cell(row=p+2,column=2,value="Bank: [A/C no] · IFSC: [ifsc] · [bank name]")
ws.cell(row=p+3,column=2,value="Terms: 50% advance, 50% on delivery. Thank you!").font=Font(italic=True,size=9,color=GREY)
c=ws.cell(row=p+5,column=2,value="💡 Edit only the YELLOW cells — totals calculate automatically."); c.font=Font(italic=True,size=9,color=TEAL)
wb.save(DIR+"Invoice-Template.xlsx"); print("✓ Invoice")

# ============ 5) CONTENT PLANNER ============
wb=Workbook(); ws=wb.active; ws.title="30-Day Planner"; ws.sheet_view.showGridLines=False
cols=["Day","Date","Platform","Content Type","Hook / Title","Caption Idea","CTA","Hashtags","Status"]
for i,w in enumerate([6,12,12,15,30,34,16,24,12],1): ws.column_dimensions[get_column_letter(i)].width=w
ws.merge_cells("A1:I1"); t=ws["A1"]; t.value="📅  30-DAY CONTENT PLANNER"; t.font=Font(bold=True,size=18,color="FFFFFF"); t.fill=navy; t.alignment=Alignment(horizontal="left",vertical="center"); ws.row_dimensions[1].height=40
ws.merge_cells("A2:I2"); s=ws["A2"]; s.value="Client: [name]    ·    Month: [month]    ·    Goal: [followers / leads / sales]    —    NextGen AI Lab · OptimityFX"; s.font=Font(italic=True,size=9,color=GREY)
hr=4
for i,h in enumerate(cols,1):
    c=ws.cell(row=hr,column=i,value=h); c.fill=navy; c.font=white; c.alignment=wrap; c.border=bd
ws.row_dimensions[hr].height=26
types=["Reel","Carousel","Story","Reel","Static Post","Reel","Carousel"]
typecolor={"Reel":"E8FBF4","Carousel":"EEF3FF","Story":"FFF1F5","Static Post":"FFF6D6"}
for d in range(1,31):
    r=hr+d; ct=types[(d-1)%7]
    ws.cell(row=r,column=1,value=d).alignment=ctr
    tcell=ws.cell(row=r,column=4,value=ct); tcell.fill=PatternFill("solid",fgColor=typecolor[ct]); tcell.alignment=ctr; tcell.font=Font(bold=True,size=9,color=INK)
    for col in range(1,10):
        cc=ws.cell(row=r,column=col); cc.border=bd
        if col in (2,5,6,7,8): cc.fill=yel
        elif col==9: cc.value="Planned"; cc.alignment=ctr; cc.font=Font(size=9,color=GREY)
        elif col in (1,3) and r%2==0: cc.fill=zebra
ws.cell(row=hr+32,column=1,value="💡 Edit YELLOW cells. Mix education / entertainment / promo in a 3:1 ratio.").font=Font(italic=True,size=9,color=TEAL)
ws.freeze_panes="A5"
wb.save(DIR+"Content-Planner.xlsx"); print("✓ Content Planner")
print("ALL 5 BUILT")
