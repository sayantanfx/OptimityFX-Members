#!/usr/bin/env python3
"""Generate editable XLSX (invoice, content planner) + DOCX (service menu, proposal) + Notion CSV."""
import csv
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DIR = "/Users/user/Claude/eBook/kit-resources/"

# ---------- shared styles ----------
ACCENT = "00B383"; DARK = "1C2138"; YELLOW = "FFF4C2"; GREY = "F2F4F8"
hdr_fill = PatternFill("solid", fgColor=DARK)
acc_fill = PatternFill("solid", fgColor=ACCENT)
in_fill  = PatternFill("solid", fgColor=YELLOW)
grey_fill= PatternFill("solid", fgColor=GREY)
white = Font(color="FFFFFF", bold=True)
bold = Font(bold=True)
thin = Side(style="thin", color="D0D5E2")
border = Border(left=thin, right=thin, top=thin, bottom=thin)

# ========== 1) INVOICE TEMPLATE ==========
wb = Workbook(); ws = wb.active; ws.title = "Invoice"
ws.sheet_view.showGridLines = False
widths = [4,26,12,14,16]
for i,w in enumerate(widths,1): ws.column_dimensions[get_column_letter(i)].width = w

ws["B2"] = "INVOICE"; ws["B2"].font = Font(bold=True, size=24, color=ACCENT)
ws["B3"] = "[Your Name / Brand]"; ws["B3"].font = bold
ws["B4"] = "[Your address / city]"
ws["B5"] = "[your@email.com] · [phone]"
ws["B6"] = "GSTIN: [if applicable]"

ws["D2"]="Invoice #:"; ws["E2"]="[INV-001]"
ws["D3"]="Date:";      ws["E3"]="[DD/MM/YYYY]"
ws["D4"]="Due date:";  ws["E4"]="[DD/MM/YYYY]"
for c in ("D2","D3","D4"): ws[c].font=bold; ws[c].alignment=Alignment(horizontal="right")
for c in ("E2","E3","E4"): ws[c].fill=in_fill; ws[c].border=border

ws["B8"]="BILL TO:"; ws["B8"].font=Font(bold=True,color=ACCENT)
ws["B9"]="[Client name / business]"; ws["B9"].fill=in_fill; ws["B9"].border=border
ws["B10"]="[Client email / phone]"; ws["B10"].fill=in_fill; ws["B10"].border=border

# table header
hr=12
heads=["#","Description","Qty","Rate (₹)","Amount (₹)"]
for i,h in enumerate(heads,1):
    c=ws.cell(row=hr,column=i,value=h); c.fill=hdr_fill; c.font=white
    c.alignment=Alignment(horizontal="center"); c.border=border
# item rows with formula
first=hr+1; last=first+5
for r in range(first,last):
    ws.cell(row=r,column=1,value=r-hr).border=border
    d=ws.cell(row=r,column=2,value="[Service description]"); d.fill=in_fill; d.border=border
    q=ws.cell(row=r,column=3,value=1); q.fill=in_fill; q.border=border; q.alignment=Alignment(horizontal="center")
    rt=ws.cell(row=r,column=4); rt.fill=in_fill; rt.border=border; rt.number_format='#,##0'
    am=ws.cell(row=r,column=5,value=f"=C{r}*D{r}"); am.border=border; am.number_format='#,##0'
# only first row gets sample rate
ws.cell(row=first,column=4,value=0)

sub=last;
ws.cell(row=sub,column=4,value="Subtotal").font=bold
sc=ws.cell(row=sub,column=5,value=f"=SUM(E{first}:E{last-1})"); sc.number_format='#,##0'; sc.font=bold
gstr=sub+1
ws.cell(row=gstr,column=2,value="GST % (edit, set 0 if none)").font=Font(italic=True,size=9,color="6B7299")
gcell=ws.cell(row=gstr,column=3,value=18); gcell.fill=in_fill; gcell.border=border; gcell.alignment=Alignment(horizontal="center")
ws.cell(row=gstr,column=4,value="GST amount").font=bold
ga=ws.cell(row=gstr,column=5,value=f"=E{sub}*C{gstr}/100"); ga.number_format='#,##0'
tot=gstr+1
tl=ws.cell(row=tot,column=4,value="TOTAL"); tl.font=white; tl.fill=acc_fill; tl.alignment=Alignment(horizontal="right")
tv=ws.cell(row=tot,column=5,value=f"=E{sub}+E{gstr}"); tv.font=white; tv.fill=acc_fill; tv.number_format='#,##0'

pay=tot+2
ws.cell(row=pay,column=2,value="PAYMENT DETAILS").font=Font(bold=True,color=ACCENT)
ws.cell(row=pay+1,column=2,value="UPI: [your-upi@bank]")
ws.cell(row=pay+2,column=2,value="Bank: [A/C no] · IFSC: [ifsc] · [bank name]")
ws.cell(row=pay+3,column=2,value="Terms: 50% advance, 50% on delivery. Thank you!").font=Font(italic=True,size=9,color="6B7299")
ws.cell(row=pay+5,column=2,value="💡 Edit only the yellow cells — totals calculate automatically.").font=Font(italic=True,size=9,color=ACCENT)
wb.save(DIR+"Invoice-Template.xlsx")
print("✓ Invoice-Template.xlsx")

# ========== 2) CONTENT PLANNER (30-day) ==========
wb2=Workbook(); ws=wb2.active; ws.title="30-Day Planner"
ws.sheet_view.showGridLines=False
cols=["Day","Date","Platform","Content Type","Hook / Title","Caption Idea","CTA","Hashtags","Status"]
cw=[6,12,12,14,28,34,18,24,12]
for i,w in enumerate(cw,1): ws.column_dimensions[get_column_letter(i)].width=w
ws["A1"]="30-DAY CONTENT PLANNER"; ws["A1"].font=Font(bold=True,size=16,color=ACCENT)
ws["A2"]="Client: [name]   ·   Month: [month]   ·   Goal: [followers / leads / sales]"; ws["A2"].font=Font(italic=True,size=9,color="6B7299")
hr=4
for i,h in enumerate(cols,1):
    c=ws.cell(row=hr,column=i,value=h); c.fill=hdr_fill; c.font=white
    c.alignment=Alignment(horizontal="center",wrap_text=True); c.border=border
types=["Reel","Carousel","Story","Reel","Static Post","Reel","Carousel"]
for d in range(1,31):
    r=hr+d
    ws.cell(row=r,column=1,value=d).alignment=Alignment(horizontal="center")
    ws.cell(row=r,column=4,value=types[(d-1)%7])
    for col in range(1,10):
        cc=ws.cell(row=r,column=col); cc.border=border
        if col in (2,5,6,7,8): cc.fill=in_fill
        if col==9: cc.value="Planned"
    if d%2==0:
        for col in range(1,10):
            if ws.cell(row=r,column=col).fill.fgColor.rgb in ("00000000",None):
                ws.cell(row=r,column=col).fill=grey_fill
ws.cell(row=hr+32,column=1,value="💡 Edit yellow cells. Mix education/entertainment/promo in a 3:1 ratio.").font=Font(italic=True,size=9,color=ACCENT)
ws.freeze_panes="A5"
wb2.save(DIR+"Content-Planner.xlsx")
print("✓ Content-Planner.xlsx")

# ========== 3) CONTENT PLANNER NOTION CSV ==========
with open(DIR+"Content-Planner-Notion.csv","w",newline="",encoding="utf-8") as f:
    w=csv.writer(f); w.writerow(cols)
    for d in range(1,31):
        w.writerow([d,"","Instagram",types[(d-1)%7],"","","","","Planned"])
print("✓ Content-Planner-Notion.csv")

# ========== 4) DOCX helper ==========
def style_doc(doc):
    n=doc.styles["Normal"]; n.font.name="Calibri"; n.font.size=Pt(11)

def h(doc,text,size=16,color=ACCENT):
    p=doc.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(size)
    r.font.color.rgb=RGBColor.from_string(color); return p

# ---- Service Menu DOCX ----
doc=Document(); style_doc(doc)
h(doc,"[Your Name / Brand] — Service Menu",20)
p=doc.add_paragraph(); r=p.add_run("AI-powered [service category] for [target clients]. Fast turnaround, fair pricing, no fluff.")
r.italic=True; r.font.color.rgb=RGBColor.from_string("6B7299")
doc.add_paragraph("[your@email.com]  ·  [WhatsApp]  ·  [social handle]")
for name,price,perfect,items in [
    ("🟢 Starter Package","₹[price]","businesses just getting started",["[Deliverable 1]","[Deliverable 2]","[Deliverable 3]","Turnaround: [X days]"]),
    ("🔵 Growth Package (most popular)","₹[price] /month","businesses that want consistent results",["Everything in Starter, plus:","[Deliverable 4]","[Deliverable 5]","Monthly report","Priority support"]),
    ("🟣 Pro Package","₹[price] /month","businesses ready to scale",["Everything in Growth, plus:","[Deliverable 6]","[Deliverable 7]","[Add-on]"]),
]:
    h(doc,name,14)
    pp=doc.add_paragraph(); rr=pp.add_run(f"Price: {price}  ·  Perfect for: {perfect}"); rr.bold=True
    for it in items: doc.add_paragraph(it,style="List Bullet")
h(doc,"Add-Ons",14)
t=doc.add_table(rows=1,cols=2); t.style="Light Grid Accent 1"
t.rows[0].cells[0].text="Add-on"; t.rows[0].cells[1].text="Price"
for a,p_ in [("[Extra service 1]","₹[xxx]"),("[Extra service 2]","₹[xxx]"),("Rush delivery (24h)","+[xx]%")]:
    row=t.add_row().cells; row[0].text=a; row[1].text=p_
h(doc,"How It Works",14)
for s in ["1. You pick a package.","2. 50% advance to start (UPI / bank).","3. I deliver within the timeline.","4. 50% on delivery. Done!"]:
    doc.add_paragraph(s)
doc.add_paragraph()
fp=doc.add_paragraph(); fr=fp.add_run("Ready to start? Message me on [WhatsApp/Instagram] or email [your@email.com]."); fr.bold=True
ft=doc.add_paragraph(); ftr=ft.add_run("Built with NextGen AI Lab · OptimityFX"); ftr.italic=True; ftr.font.size=Pt(9); ftr.font.color.rgb=RGBColor.from_string("6B7299")
doc.save(DIR+"02-Service-Menu.docx")
print("✓ 02-Service-Menu.docx")

# ---- Proposal DOCX ----
doc=Document(); style_doc(doc)
h(doc,"Project Proposal",20)
for k,v in [("Prepared for","[Client Name / Business]"),("Prepared by","[Your Name / Brand]"),("Date","[Date]")]:
    pp=doc.add_paragraph(); rr=pp.add_run(f"{k}: "); rr.bold=True; pp.add_run(v)
h(doc,"1. What You Need",13)
doc.add_paragraph("[Describe the client's problem in 1–2 lines. Show you understood them.]")
h(doc,"2. What I'll Deliver",13)
for it in ["[Deliverable 1 — with quantity]","[Deliverable 2]","[Deliverable 3]","[Deliverable 4]"]:
    doc.add_paragraph(it,style="List Bullet")
h(doc,"3. Timeline",13)
t=doc.add_table(rows=1,cols=3); t.style="Light Grid Accent 1"
for i,x in enumerate(["Phase","What happens","When"]): t.rows[0].cells[i].text=x
for ph,wha,wn in [("Kickoff","Brief + access + samples","Day 1"),("Draft","First batch for review","Day [x]"),("Delivery","Final, revised","Day [x]")]:
    c=t.add_row().cells; c[0].text=ph; c[1].text=wha; c[2].text=wn
h(doc,"4. Investment",13)
t=doc.add_table(rows=1,cols=2); t.style="Light Grid Accent 1"
t.rows[0].cells[0].text="Item"; t.rows[0].cells[1].text="Price"
for a,b in [("[Package name]","₹[price]"),("[Add-on, optional]","₹[price]"),("Total","₹[total]")]:
    c=t.add_row().cells; c[0].text=a; c[1].text=b
doc.add_paragraph("Payment terms: 50% advance to begin, 50% on delivery (UPI / bank).")
h(doc,"5. Why Me",13)
for it in ["[Reason 1]","[Reason 2 — link 2 samples]","[Reason 3]"]:
    doc.add_paragraph(it,style="List Bullet")
doc.add_paragraph()
fp=doc.add_paragraph(); fr=fp.add_run('Ready to start? Reply "yes" and I\'ll send the payment link. Valid until [date].'); fr.bold=True
doc.add_paragraph("[your@email.com] · [WhatsApp]")
ft=doc.add_paragraph(); ftr=ft.add_run("Built with NextGen AI Lab · OptimityFX"); ftr.italic=True; ftr.font.size=Pt(9); ftr.font.color.rgb=RGBColor.from_string("6B7299")
doc.save(DIR+"05-Proposal-Template.docx")
print("✓ 05-Proposal-Template.docx")

# ---- Portfolio DOCX ----
doc=Document(); style_doc(doc)
h(doc,"[Your Name] — [Your Service] for [Type of Client]",18)
p=doc.add_paragraph(); r=p.add_run("Helping [target audience] get [the result you deliver] using AI-powered systems.")
r.italic=True; r.font.color.rgb=RGBColor.from_string("6B7299")
doc.add_paragraph("[your@email.com]  ·  [WhatsApp]  ·  [Instagram/LinkedIn]")
h(doc,"About Me",13)
doc.add_paragraph("I help [type of business] with [your service]. I use AI tools to deliver fast, high-quality work at a fair price. [2–3 honest lines.]")
h(doc,"What I Do",13)
for it in ["[Service 1]","[Service 2]","[Service 3]"]: doc.add_paragraph(it,style="List Bullet")
h(doc,"Work Samples",13)
doc.add_paragraph("Replace with your 3 best samples (screenshots/links). New? Make 3 free samples first — that IS your portfolio.")
for i in (1,2,3):
    pp=doc.add_paragraph(); rr=pp.add_run(f"Sample {i} — [Client/Project]"); rr.bold=True
    doc.add_paragraph("What I did: [1 line]  ·  Result: [1 line]")
h(doc,"Packages",13)
t=doc.add_table(rows=1,cols=3); t.style="Light Grid Accent 1"
for i,x in enumerate(["Package","What's included","Price"]): t.rows[0].cells[i].text=x
for a,b,c_ in [("Starter","[deliverables]","₹[xxx]"),("Growth","[deliverables]","₹[xxx]"),("Pro","[deliverables]","₹[xxx]")]:
    c=t.add_row().cells; c[0].text=a; c[1].text=b; c[2].text=c_
doc.add_paragraph()
fp=doc.add_paragraph(); fr=fp.add_run("Let's work together — message me on [WhatsApp/Instagram] or email [your@email.com]."); fr.bold=True
ft=doc.add_paragraph(); ftr=ft.add_run("Powered by NextGen AI Lab · OptimityFX"); ftr.italic=True; ftr.font.size=Pt(9); ftr.font.color.rgb=RGBColor.from_string("6B7299")
doc.save(DIR+"01-Portfolio-Template.docx")
print("✓ 01-Portfolio-Template.docx")
print("ALL FILES BUILT")
